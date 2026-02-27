"""
Legacy Document Import Service.

Processes legacy architecture documents (PDF / DOCX) and uses AI to extract
structured pattern entities (AB, ABB, SBB, Technology, PBC) for review.

Pipeline:
  1. Extract text or images from the uploaded document.
  2. AI Phase 1 — Document overview (type, sections, metadata).
  3. AI Phase 2 — Entity extraction (patterns, technologies, capabilities).
  4. Phase 3 — Cross-reference against existing catalog via vector search.
  5. Compile and save analysis report.
"""

import json
import logging
import os
from typing import AsyncIterator, Callable, Optional

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Document text / image extraction
# ---------------------------------------------------------------------------

ALLOWED_EXTENSIONS = {".pdf", ".docx"}
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB
MAX_PAGES = 50  # cap for scanned PDFs → limits vision API token cost


def extract_content(file_path: str) -> dict:
    """Detect document type and extract content.

    Returns:
        {
            "input_type": "text" | "images",
            "document_format": "digital_pdf" | "scanned_pdf" | "docx",
            "content": str (for text) | list[bytes] (for images),
            "page_count": int,
        }
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return _extract_from_pdf(file_path)
    elif ext == ".docx":
        return _extract_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


ANTHROPIC_IMAGE_MAX_BYTES = 4_500_000  # stay safely under 5 MB API limit


def _extract_from_pdf(pdf_path: str) -> dict:
    """Extract from PDF — auto-detects scanned vs digital."""
    import fitz  # pymupdf

    doc = fitz.open(pdf_path)
    page_count = len(doc)

    # Try extracting text from all pages
    texts = []
    for page in doc:
        texts.append(page.get_text("text").strip())
    total_text = "\n\n".join(t for t in texts if t)

    # Heuristic: if total text is very short relative to page count, it's scanned images
    avg_chars_per_page = len(total_text) / max(page_count, 1)

    if avg_chars_per_page > 100:
        # Digital PDF with extractable text
        doc.close()
        return {
            "input_type": "text",
            "document_format": "digital_pdf",
            "content": total_text,
            "page_count": page_count,
        }
    else:
        # Scanned images — convert pages to PNG, auto-scale to stay under API limit
        images = []
        for page_num in range(min(page_count, MAX_PAGES)):
            page = doc[page_num]
            png_bytes = _render_page_under_limit(page)
            images.append(png_bytes)
        doc.close()
        return {
            "input_type": "images",
            "document_format": "scanned_pdf",
            "content": images,
            "page_count": page_count,
        }


def _render_page_under_limit(page, initial_dpi: int = 200) -> bytes:
    """Render a PDF page to PNG, reducing DPI until it fits under the API limit.

    Starts at *initial_dpi* (default 200) and steps down by 25 DPI each attempt.
    Minimum DPI is 100 to keep text legible for vision models.
    """
    import fitz  # pymupdf

    dpi = initial_dpi
    while dpi >= 100:
        pix = page.get_pixmap(dpi=dpi)
        png_bytes = pix.tobytes("png")
        if len(png_bytes) <= ANTHROPIC_IMAGE_MAX_BYTES:
            if dpi < initial_dpi:
                logger.info(
                    "Page rendered at %d DPI (%d KB) to stay within API limit",
                    dpi, len(png_bytes) // 1024,
                )
            return png_bytes
        logger.debug(
            "Page at %d DPI = %d KB (exceeds %d KB limit), reducing DPI...",
            dpi, len(png_bytes) // 1024, ANTHROPIC_IMAGE_MAX_BYTES // 1024,
        )
        dpi -= 25

    # Absolute fallback: render at 100 DPI and return whatever we get
    pix = page.get_pixmap(dpi=100)
    png_bytes = pix.tobytes("png")
    logger.warning(
        "Page at minimum DPI 100 is still %d KB — sending anyway",
        len(png_bytes) // 1024,
    )
    return png_bytes


def _extract_from_docx(docx_path: str) -> dict:
    """Extract from Word document — paragraphs + tables as structured text."""
    from docx import Document

    doc = Document(docx_path)
    parts = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            # Paragraph
            text = element.text.strip() if hasattr(element, "text") and element.text else ""
            if text:
                parts.append(text)

        elif tag == "tbl":
            # Table — render as markdown
            table_rows = []
            for tr in element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr"):
                cells = []
                for tc in tr.findall(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc"
                ):
                    cell_text = ""
                    for p in tc.findall(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"
                    ):
                        for run in p.findall(
                            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r"
                        ):
                            t = run.find(
                                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"
                            )
                            if t is not None and t.text:
                                cell_text += t.text
                    cells.append(cell_text.strip())
                table_rows.append(" | ".join(cells))
            if table_rows:
                parts.append("\n".join(table_rows))

    # Fallback to simple paragraph extraction if XML parsing yielded nothing
    if not parts:
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append(" | ".join(cells))

    content = "\n\n".join(parts)
    # Rough page estimate (Word docs don't have real page counts)
    estimated_pages = max(1, len(content) // 3000)

    return {
        "input_type": "text",
        "document_format": "docx",
        "content": content,
        "page_count": estimated_pages,
    }


# ---------------------------------------------------------------------------
# AI Analysis Pipeline
# ---------------------------------------------------------------------------

ProgressCallback = Callable[[str, int, int, str], None]


async def run_analysis_pipeline(
    file_path: str,
    filename: str,
    user_email: str,
    db,
    progress: Optional[ProgressCallback] = None,
    provider_name: Optional[str] = None,
    model: Optional[str] = None,
) -> dict:
    """Full 3-phase analysis pipeline. Returns saved analysis report dict.

    Args:
        file_path: Path to the uploaded document.
        filename: Original filename.
        user_email: Email of the admin who uploaded.
        db: Neo4jService instance.
        progress: Optional callback(stage, step, total, message).
        provider_name: LLM provider override.
        model: LLM model override.
    """
    from services.llm import get_provider
    from services.prompt_service import get_merged_prompts
    from services.embedding_service import EmbeddingService

    def _progress(stage, step, total, message):
        if progress:
            progress(stage, step, total, message)

    # --- Extract document content ---
    _progress("extraction", 0, 3, f"Extracting content from {filename}...")
    extraction = extract_content(file_path)
    _progress(
        "extraction", 0, 3,
        f"Detected {extraction['document_format']} ({extraction['page_count']} pages)"
    )

    # --- Get LLM provider and prompts ---
    provider = get_provider(provider_name)
    prompts = get_merged_prompts()
    li_prompts = prompts.get("legacy_import", {})
    system_prompt = li_prompts.get("system", "")

    used_model = model or provider.default_model

    # --- Phase 1: Document Overview ---
    _progress("overview", 1, 4, "Analyzing document structure...")

    if extraction["input_type"] == "text":
        # Send structural preview spanning the full document
        text_preview = _build_structural_preview(extraction["content"], extraction["page_count"])
        content_block = f"## Document Content (structural preview, {extraction['page_count']} pages total)\n\n{text_preview}"
        overview_prompt = li_prompts.get("overview", "").replace("{content_block}", content_block)
        overview_result = await provider.generate(system_prompt, overview_prompt, used_model)
    else:
        # Send up to 10 page images for overview (was 3 — need full document visibility)
        max_overview_pages = min(len(extraction["content"]), 10)
        preview_images = extraction["content"][:max_overview_pages]
        content_block = f"The document pages are provided as images above ({max_overview_pages} of {extraction['page_count']} pages). Analyze them."
        overview_prompt = li_prompts.get("overview", "").replace("{content_block}", content_block)
        overview_result = await provider.generate_with_images(
            system_prompt, overview_prompt, preview_images, used_model
        )

    overview = _parse_json_response(overview_result["content"], default={
        "document_type": "other",
        "confidence": "LOW",
        "title": filename,
        "sections": [],
        "summary": "Could not parse overview.",
    })
    _progress("overview", 1, 4, f"Document identified as: {overview.get('document_type', 'unknown')}")

    # --- Phase 2: Entity Extraction ---
    _progress("entity_extraction", 2, 4, "Extracting entities from document...")

    # Build catalog context (categories + existing pattern/tech/PBC summaries)
    catalog_context = _build_catalog_context(db)

    overview_json = json.dumps(overview, indent=2)

    if extraction["input_type"] == "text":
        # Send full text (capped at ~30K chars to stay within context window)
        full_text = extraction["content"][:30000]
        content_block = f"## Full Document Content\n\n{full_text}"
        extract_prompt = (
            li_prompts.get("extract", "")
            .replace("{content_block}", content_block)
            .replace("{overview_json}", overview_json)
            .replace("{catalog_context}", catalog_context)
        )
        extract_result = await provider.generate(system_prompt, extract_prompt, used_model)
    else:
        # Send all page images
        content_block = "The full document pages are provided as images above. Extract all entities."
        extract_prompt = (
            li_prompts.get("extract", "")
            .replace("{content_block}", content_block)
            .replace("{overview_json}", overview_json)
            .replace("{catalog_context}", catalog_context)
        )
        extract_result = await provider.generate_with_images(
            system_prompt, extract_prompt, extraction["content"], used_model
        )

    entities = _parse_json_response(extract_result["content"], default={
        "primary_pattern": None,
        "sub_patterns": [],
        "technologies": [],
        "business_capabilities": [],
        "relationships": [],
        "skipped_content": [],
    })
    _progress("entity_extraction", 2, 4, f"Extracted {_count_entities(entities)} entities")

    # --- Phase 2.5: Quality Gate ---
    # Build content block for judge/critic (reuse the extraction content)
    if extraction["input_type"] == "text":
        qg_content_block = f"## Full Document Content\n\n{extraction['content'][:30000]}"
    else:
        qg_content_block = "The full document pages were provided as images during extraction."

    entities, quality_scores = await _quality_gate(
        overview=overview,
        entities=entities,
        content_block=qg_content_block,
        provider=provider,
        system_prompt=system_prompt,
        li_prompts=li_prompts,
        model=used_model,
        progress_cb=_progress,
    )

    # --- Phase 3: Cross-Reference ---
    _progress("cross_reference", 4, 4, "Cross-referencing with existing catalog...")
    cross_refs = _cross_reference_entities(entities, db)
    _progress("cross_reference", 4, 4, "Cross-referencing complete")

    # --- Compile and save report ---
    summary = _build_summary(entities, cross_refs, quality_scores)

    report_data = {
        "title": f"Analysis: {overview.get('title', filename)}",
        "filename": filename,
        "document_type": extraction["document_format"],
        "page_count": extraction["page_count"],
        "overview": overview,
        "entities": entities,
        "cross_references": cross_refs,
        "summary": summary,
        "provider": provider.name,
        "model": used_model,
        "created_by": user_email,
    }

    saved = db.save_legacy_analysis(report_data)
    _progress("complete", 3, 3, f"Analysis saved as {saved['id']}")

    return saved


# ---------------------------------------------------------------------------
# Chat Refinement
# ---------------------------------------------------------------------------

async def chat_with_analysis(
    analysis: dict,
    message: str,
    db,
    provider_name: Optional[str] = None,
    model: Optional[str] = None,
) -> dict:
    """Send a chat message to refine an analysis.

    Returns:
        {"response": str, "entity_updates": list | None}
    """
    from services.llm import get_provider
    from services.prompt_service import get_merged_prompts

    provider = get_provider(provider_name)
    prompts = get_merged_prompts()
    li_prompts = prompts.get("legacy_import", {})

    system_prompt = li_prompts.get("system", "")
    chat_template = li_prompts.get("chat", "")

    # Build conversation history
    messages = analysis.get("messages_json", [])
    if isinstance(messages, str):
        try:
            messages = json.loads(messages)
        except Exception:
            messages = []

    history_text = ""
    for msg in messages:
        role = msg.get("role", "user")
        content = msg.get("content", "")
        history_text += f"\n**{role.title()}**: {content}\n"

    chat_prompt = (
        chat_template
        .replace("{overview_json}", json.dumps(analysis.get("overview_json", {}), indent=2))
        .replace("{entities_json}", json.dumps(analysis.get("entities_json", {}), indent=2))
        .replace("{cross_references_json}", json.dumps(analysis.get("cross_references_json", {}), indent=2))
        .replace("{conversation_history}", history_text or "No previous messages.")
        .replace("{message}", message)
    )

    used_model = model or provider.default_model
    result = await provider.generate(system_prompt, chat_prompt, used_model)
    response_text = result["content"]

    # Check for entity update block
    entity_updates = None
    if "---ENTITY_UPDATE---" in response_text:
        parts = response_text.split("---ENTITY_UPDATE---", 1)
        response_text = parts[0].strip()
        update_block = parts[1].strip()
        # Extract JSON from code fence
        if "```json" in update_block:
            json_str = update_block.split("```json", 1)[1].split("```", 1)[0].strip()
        elif "```" in update_block:
            json_str = update_block.split("```", 1)[1].split("```", 1)[0].strip()
        else:
            json_str = update_block
        try:
            entity_updates = json.loads(json_str)
        except json.JSONDecodeError:
            logger.warning("Failed to parse entity update JSON from chat response")

    # Update chat history
    messages.append({"role": "user", "content": message})
    messages.append({"role": "assistant", "content": response_text})

    db.update_legacy_analysis_chat(
        analysis["id"],
        json.dumps(messages),
        len(messages),
    )

    # If entity updates provided, apply them
    if entity_updates and entity_updates.get("updates"):
        entities = analysis.get("entities_json", {})
        if isinstance(entities, str):
            try:
                entities = json.loads(entities)
            except Exception:
                entities = {}
        entities = _apply_entity_updates(entities, entity_updates["updates"])
        # Recount summary
        cross_refs = analysis.get("cross_references_json", {})
        if isinstance(cross_refs, str):
            try:
                cross_refs = json.loads(cross_refs)
            except Exception:
                cross_refs = {}
        summary = _build_summary(entities, cross_refs)
        db.update_legacy_analysis_entities(
            analysis["id"],
            json.dumps(entities),
            json.dumps(summary),
        )

    return {
        "response": response_text,
        "entity_updates": entity_updates,
    }


async def chat_stream_tokens(
    analysis: dict,
    message: str,
    provider_name: Optional[str] = None,
    model: Optional[str] = None,
) -> AsyncIterator[str]:
    """Stream chat response tokens. Caller should collect full text and
    call finalize_chat_stream() after exhausting this generator."""
    from services.llm import get_provider
    from services.prompt_service import get_merged_prompts

    provider = get_provider(provider_name)
    prompts = get_merged_prompts()
    li_prompts = prompts.get("legacy_import", {})

    system_prompt = li_prompts.get("system", "")
    chat_template = li_prompts.get("chat", "")

    # Build conversation history
    messages = analysis.get("messages_json", [])
    if isinstance(messages, str):
        try:
            messages = json.loads(messages)
        except Exception:
            messages = []

    history_text = ""
    for msg in messages:
        role = msg.get("role", "user")
        content = msg.get("content", "")
        history_text += f"\n**{role.title()}**: {content}\n"

    chat_prompt = (
        chat_template
        .replace("{overview_json}", json.dumps(analysis.get("overview_json", {}), indent=2))
        .replace("{entities_json}", json.dumps(analysis.get("entities_json", {}), indent=2))
        .replace("{cross_references_json}", json.dumps(analysis.get("cross_references_json", {}), indent=2))
        .replace("{conversation_history}", history_text or "No previous messages.")
        .replace("{message}", message)
    )

    used_model = model or provider.default_model
    async for chunk in provider.generate_stream(system_prompt, chat_prompt, used_model):
        yield chunk


def finalize_chat_stream(
    analysis: dict,
    message: str,
    full_response: str,
    db,
):
    """Persist chat history and apply entity updates after streaming completes."""
    response_text = full_response

    # Check for entity update block
    entity_updates = None
    if "---ENTITY_UPDATE---" in response_text:
        parts = response_text.split("---ENTITY_UPDATE---", 1)
        response_text = parts[0].strip()
        update_block = parts[1].strip()
        if "```json" in update_block:
            json_str = update_block.split("```json", 1)[1].split("```", 1)[0].strip()
        elif "```" in update_block:
            json_str = update_block.split("```", 1)[1].split("```", 1)[0].strip()
        else:
            json_str = update_block
        try:
            entity_updates = json.loads(json_str)
        except json.JSONDecodeError:
            logger.warning("Failed to parse entity update JSON from chat response")

    # Update chat history
    messages = analysis.get("messages_json", [])
    if isinstance(messages, str):
        try:
            messages = json.loads(messages)
        except Exception:
            messages = []
    messages.append({"role": "user", "content": message})
    messages.append({"role": "assistant", "content": response_text})

    db.update_legacy_analysis_chat(
        analysis["id"],
        json.dumps(messages),
        len(messages),
    )

    # If entity updates provided, apply them
    if entity_updates and entity_updates.get("updates"):
        entities = analysis.get("entities_json", {})
        if isinstance(entities, str):
            try:
                entities = json.loads(entities)
            except Exception:
                entities = {}
        entities = _apply_entity_updates(entities, entity_updates["updates"])
        cross_refs = analysis.get("cross_references_json", {})
        if isinstance(cross_refs, str):
            try:
                cross_refs = json.loads(cross_refs)
            except Exception:
                cross_refs = {}
        summary = _build_summary(entities, cross_refs)
        db.update_legacy_analysis_entities(
            analysis["id"],
            json.dumps(entities),
            json.dumps(summary),
        )


# ---------------------------------------------------------------------------
# Helper Functions
# ---------------------------------------------------------------------------

def _build_structural_preview(full_text: str, page_count: int, max_chars: int = 8000) -> str:
    """Build a structural preview spanning the full document.

    Strategy:
      - First ~2000 chars verbatim (captures title, metadata, headers)
      - For multi-page docs: sample ~500 chars around each estimated page boundary
      - Add ``--- Page N (approx) ---`` markers so the LLM sees the full span
      - Cap total output at *max_chars*

    Works for any document length — 3-page or 50-page.
    """
    if len(full_text) <= max_chars:
        return full_text  # Short doc — return everything

    parts: list[str] = []
    header_size = 2000
    sample_size = 500

    # --- Header: first ~2000 chars ---
    parts.append(full_text[:header_size])

    # --- Estimate page boundaries and sample around each ---
    if page_count <= 1:
        # Single page but long text — just take beginning and end
        tail = full_text[-sample_size:]
        parts.append(f"\n\n--- End of document (approx) ---\n{tail}")
    else:
        chars_per_page = len(full_text) / page_count

        # Sample pages: skip page 1 (covered by header), sample the rest
        # For very long docs, sample every Nth page to stay within budget
        pages_budget = (max_chars - header_size) // (sample_size + 60)  # 60 chars for marker
        if page_count - 1 > pages_budget:
            step = max(1, (page_count - 1) // pages_budget)
            sample_pages = list(range(2, page_count + 1, step))
        else:
            sample_pages = list(range(2, page_count + 1))

        for page_num in sample_pages:
            boundary = int((page_num - 1) * chars_per_page)
            start = max(header_size, boundary - sample_size // 2)
            end = min(len(full_text), start + sample_size)
            if start >= len(full_text):
                break
            snippet = full_text[start:end]
            parts.append(f"\n\n--- Page {page_num} (approx) ---\n{snippet}")

    preview = "".join(parts)
    if len(preview) > max_chars:
        preview = preview[:max_chars] + "\n[...truncated...]"
    return preview


def _parse_json_response(text: str, default: dict) -> dict:
    """Parse JSON from LLM response, handling code fences and extra text."""
    text = text.strip()
    # Strip markdown code fences
    if text.startswith("```json"):
        text = text[7:]
    elif text.startswith("```"):
        text = text[3:]
    if text.endswith("```"):
        text = text[:-3]
    text = text.strip()

    # Try direct parse
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    # Try to find JSON object in the text
    start = text.find("{")
    if start >= 0:
        # Find matching closing brace
        depth = 0
        for i in range(start, len(text)):
            if text[i] == "{":
                depth += 1
            elif text[i] == "}":
                depth -= 1
                if depth == 0:
                    try:
                        return json.loads(text[start : i + 1])
                    except json.JSONDecodeError:
                        break

    logger.warning("Failed to parse JSON from LLM response, using default")
    return default


def _build_catalog_context(db) -> str:
    """Build a summary of existing catalog for LLM context."""
    lines = []

    # Categories
    try:
        categories = db.list_categories()
        if categories:
            lines.append("## Existing Categories")
            for c in categories:
                lines.append(f"- {c['code']}: {c.get('label', c['code'])}")
            lines.append("")
    except Exception:
        pass

    # Pattern counts and examples
    try:
        patterns, total = db.list_patterns(limit=200)
        if patterns:
            type_counts = {}
            for p in patterns:
                t = p.get("type", "?")
                type_counts[t] = type_counts.get(t, 0) + 1
            lines.append(f"## Existing Patterns ({total} total)")
            for t, c in sorted(type_counts.items()):
                lines.append(f"- {t}: {c} patterns")
            lines.append("")
            # List a few names per type for reference
            for ptype in ["AB", "ABB", "SBB"]:
                typed = [p for p in patterns if p.get("type") == ptype][:5]
                if typed:
                    names = ", ".join(f"{p['id']} ({p['name']})" for p in typed)
                    lines.append(f"Sample {ptype}s: {names}")
            lines.append("")
    except Exception:
        pass

    # Technologies
    try:
        techs, total = db.list_technologies(limit=200)
        if techs:
            lines.append(f"## Existing Technologies ({total} total)")
            for t in techs[:20]:
                lines.append(f"- {t['id']}: {t['name']} ({t.get('vendor', '')}) [{t.get('category', '')}]")
            if total > 20:
                lines.append(f"... and {total - 20} more")
            lines.append("")
    except Exception:
        pass

    # PBCs
    try:
        pbcs = db.list_pbcs()
        if pbcs:
            lines.append(f"## Existing Business Capabilities ({len(pbcs)} total)")
            for p in pbcs[:10]:
                lines.append(f"- {p['id']}: {p['name']}")
            lines.append("")
    except Exception:
        pass

    # ID conventions
    lines.append("## ID Conventions")
    lines.append("- AB: AB-PAT-NNN")
    lines.append("- ABB: ABB-{CAT}-NNN (e.g., ABB-CORE-001, ABB-INTG-002)")
    lines.append("- SBB: SBB-{CAT}-NNN")
    lines.append("- Technology: kebab-case-slug (e.g., aws-bedrock, apigee)")
    lines.append("- PBC: PBC-NNN")

    return "\n".join(lines)


def _cross_reference_entities(entities: dict, db) -> dict:
    """Run vector search for each extracted entity against existing catalog."""
    from services.embedding_service import EmbeddingService

    embed_svc = EmbeddingService()
    cross_refs = {
        "patterns": [],
        "technologies": [],
        "business_capabilities": [],
    }

    if not embed_svc.available:
        logger.info("Embedding service unavailable — skipping cross-reference")
        return cross_refs

    # Cross-ref primary pattern
    primary = entities.get("primary_pattern")
    if primary:
        cross_refs["patterns"].append(
            _search_similar_pattern(embed_svc, db, primary.get("name", ""), primary)
        )

    # Cross-ref sub-patterns
    for sp in entities.get("sub_patterns", []):
        cross_refs["patterns"].append(
            _search_similar_pattern(embed_svc, db, sp.get("name", ""), sp)
        )

    # Cross-ref technologies
    for tech in entities.get("technologies", []):
        name = tech.get("name", "")
        if not name:
            continue
        try:
            query_text = f"{name} {tech.get('vendor', '')} {tech.get('description', '')}"
            embedding = embed_svc.generate_embedding(query_text)
            matches = db.vector_search_technologies(embedding, limit=3)
            match_status = "new"
            if matches:
                top_score = matches[0].get("score", 0)
                if top_score > 0.9:
                    match_status = "likely_duplicate"
                elif top_score > 0.7:
                    match_status = "similar_exists"
            cross_refs["technologies"].append({
                "name": name,
                "match_status": match_status,
                "similar_existing": [
                    {"id": m["id"], "name": m["name"], "similarity": round(m.get("score", 0), 3)}
                    for m in matches[:3]
                ],
            })
        except Exception as e:
            logger.warning(f"Cross-ref failed for technology '{name}': {e}")
            cross_refs["technologies"].append({
                "name": name, "match_status": "unknown", "similar_existing": [],
            })

    # Cross-ref business capabilities (PBCs)
    for cap in entities.get("business_capabilities", []):
        name = cap.get("name", "")
        if not name:
            continue
        try:
            query_text = f"{name} {cap.get('description', '')}"
            embedding = embed_svc.generate_embedding(query_text)
            matches = db.vector_search_pbcs(embedding, limit=3)
            match_status = "new"
            if matches:
                top_score = matches[0].get("score", 0)
                if top_score > 0.9:
                    match_status = "likely_duplicate"
                elif top_score > 0.7:
                    match_status = "similar_exists"
            cross_refs["business_capabilities"].append({
                "name": name,
                "match_status": match_status,
                "similar_existing": [
                    {"id": m["id"], "name": m["name"], "similarity": round(m.get("score", 0), 3)}
                    for m in matches[:3]
                ],
            })
        except Exception as e:
            logger.warning(f"Cross-ref failed for capability '{name}': {e}")
            cross_refs["business_capabilities"].append({
                "name": name, "match_status": "unknown", "similar_existing": [],
            })

    return cross_refs


def _search_similar_pattern(embed_svc, db, name: str, entity: dict) -> dict:
    """Vector search for a similar pattern."""
    try:
        fields = entity.get("fields", {})
        desc = fields.get("description", "") or fields.get("intent", "") or fields.get("functionality", "")
        query_text = f"{name} {desc}"
        embedding = embed_svc.generate_embedding(query_text)
        matches = db.vector_search_patterns(embedding, limit=3)
        match_status = "new"
        if matches:
            top_score = matches[0].get("score", 0)
            if top_score > 0.9:
                match_status = "likely_duplicate"
            elif top_score > 0.7:
                match_status = "similar_exists"
        return {
            "name": name,
            "match_status": match_status,
            "similar_existing": [
                {
                    "id": m["id"],
                    "name": m["name"],
                    "type": m.get("type", ""),
                    "similarity": round(m.get("score", 0), 3),
                }
                for m in matches[:3]
            ],
        }
    except Exception as e:
        logger.warning(f"Cross-ref failed for pattern '{name}': {e}")
        return {"name": name, "match_status": "unknown", "similar_existing": []}


def _count_entities(entities: dict) -> int:
    """Count total entities in extraction result."""
    count = 0
    if entities.get("primary_pattern"):
        count += 1
    count += len(entities.get("sub_patterns", []))
    count += len(entities.get("technologies", []))
    count += len(entities.get("business_capabilities", []))
    return count


async def _quality_gate(
    overview: dict,
    entities: dict,
    content_block: str,
    provider,
    system_prompt: str,
    li_prompts: dict,
    model: str,
    progress_cb: Optional[ProgressCallback] = None,
) -> tuple:
    """Run judge → critic loop to improve extraction quality.

    Returns:
        (refined_entities, quality_scores) where quality_scores is a dict with
        scores, iterations, pass status, and critique history.
    """
    MAX_ITERATIONS = 3
    PASS_THRESHOLD = 7

    def _progress(stage, step, total, message):
        if progress_cb:
            progress_cb(stage, step, total, message)

    judge_template = li_prompts.get("judge", "")
    critic_template = li_prompts.get("critic", "")

    if not judge_template or not critic_template:
        logger.warning("Judge or critic prompt not configured — skipping quality gate")
        return entities, {"skipped": True, "reason": "prompts_not_configured"}

    current_entities = entities
    quality_scores = {
        "iterations": 0,
        "passed": False,
        "scores": {},
        "critique_history": [],
    }

    for iteration in range(1, MAX_ITERATIONS + 1):
        quality_scores["iterations"] = iteration
        _progress("quality_gate", 3, 4, f"Quality review iteration {iteration}/{MAX_ITERATIONS}...")

        # --- Call Judge ---
        judge_prompt = (
            judge_template
            .replace("{overview_json}", json.dumps(overview, indent=2))
            .replace("{entities_json}", json.dumps(current_entities, indent=2))
            .replace("{content_block}", content_block)
        )

        try:
            judge_result = await provider.generate(system_prompt, judge_prompt, model)
            judge_data = _parse_json_response(judge_result["content"], default={
                "scores": {},
                "overall": 0,
                "critiques": [],
                "pass": False,
            })
        except Exception as e:
            logger.error(f"Quality gate judge call failed (iteration {iteration}): {e}")
            quality_scores["error"] = str(e)
            break

        scores = judge_data.get("scores", {})
        overall = judge_data.get("overall", 0)
        passed = judge_data.get("pass", False)
        critiques = judge_data.get("critiques", [])

        quality_scores["scores"] = scores
        quality_scores["overall"] = overall
        quality_scores["critique_history"].append({
            "iteration": iteration,
            "scores": scores,
            "overall": overall,
            "passed": passed,
            "critique_count": len(critiques),
        })

        logger.info(
            f"Quality gate iteration {iteration}: overall={overall}, "
            f"scores={scores}, pass={passed}, critiques={len(critiques)}"
        )

        if passed:
            quality_scores["passed"] = True
            _progress("quality_gate", 3, 4, f"Quality gate passed (score: {overall:.1f}/10)")
            break

        # --- Filter to only failing dimensions ---
        failing_critiques = [c for c in critiques if c.get("score", 10) < PASS_THRESHOLD]
        if not failing_critiques:
            # All dimensions OK but judge said not passed — use all critiques
            failing_critiques = critiques

        if not failing_critiques:
            # No critiques at all but not passed — nothing to fix
            quality_scores["passed"] = False
            _progress("quality_gate", 3, 4, f"Quality gate: no actionable critiques (score: {overall:.1f}/10)")
            break

        failing_dims = list(set(c.get("dimension", "?") for c in failing_critiques))
        _progress(
            "quality_gate", 3, 4,
            f"Improving {', '.join(failing_dims)} (iteration {iteration})..."
        )

        # --- Call Critic ---
        critic_prompt = (
            critic_template
            .replace("{overview_json}", json.dumps(overview, indent=2))
            .replace("{entities_json}", json.dumps(current_entities, indent=2))
            .replace("{content_block}", content_block)
            .replace("{critiques_json}", json.dumps(failing_critiques, indent=2))
        )

        try:
            critic_result = await provider.generate(system_prompt, critic_prompt, model)
            corrected = _parse_json_response(critic_result["content"], default=current_entities)
            # Validate that critic returned a sensible structure
            if corrected.get("primary_pattern") is not None or corrected.get("sub_patterns") is not None:
                current_entities = corrected
            else:
                logger.warning("Critic returned invalid structure — keeping previous entities")
        except Exception as e:
            logger.error(f"Quality gate critic call failed (iteration {iteration}): {e}")
            quality_scores["error"] = str(e)
            break
    else:
        # Exhausted all iterations without passing
        _progress(
            "quality_gate", 3, 4,
            f"Quality gate: {MAX_ITERATIONS} iterations used (final score: {quality_scores.get('overall', 0):.1f}/10)"
        )

    return current_entities, quality_scores


def _build_summary(entities: dict, cross_refs: dict, quality_scores: dict = None) -> dict:
    """Build summary statistics from entities, cross-references, and quality scores."""
    breakdown = {"AB": 0, "ABB": 0, "SBB": 0, "Technology": 0, "PBC": 0}

    primary = entities.get("primary_pattern")
    if primary:
        ptype = primary.get("suggested_type", "AB")
        if ptype in breakdown:
            breakdown[ptype] += 1

    for sp in entities.get("sub_patterns", []):
        ptype = sp.get("suggested_type", "ABB")
        if ptype in breakdown:
            breakdown[ptype] += 1

    breakdown["Technology"] = len(entities.get("technologies", []))
    breakdown["PBC"] = len(entities.get("business_capabilities", []))

    total = sum(breakdown.values())

    # Count match statuses from cross-references
    new_count = 0
    similar_count = 0
    duplicate_count = 0
    for category in ["patterns", "technologies", "business_capabilities"]:
        for ref in cross_refs.get(category, []):
            status = ref.get("match_status", "new")
            if status == "new":
                new_count += 1
            elif status == "similar_exists":
                similar_count += 1
            elif status == "likely_duplicate":
                duplicate_count += 1
            else:
                new_count += 1  # unknown → treat as new

    # Count skipped content entries
    skipped_content_count = len(entities.get("skipped_content", []))

    result = {
        "total_entities": total,
        "new_entities": new_count,
        "similar_existing": similar_count,
        "likely_duplicates": duplicate_count,
        "breakdown": breakdown,
        "skipped_content_count": skipped_content_count,
    }

    # Include quality gate scores if available
    if quality_scores and not quality_scores.get("skipped"):
        result["quality_scores"] = quality_scores

    return result


def _apply_entity_updates(entities: dict, updates: list) -> dict:
    """Apply AI-suggested entity modifications."""
    for update in updates:
        action = update.get("action")
        entity_type = update.get("entity_type")
        entity_name = update.get("entity_name", "")
        changes = update.get("changes", {})

        if action == "modify":
            if entity_type == "primary_pattern" and entities.get("primary_pattern"):
                if changes.get("fields"):
                    entities["primary_pattern"]["fields"].update(changes["fields"])
                for key in ["suggested_type", "confidence", "name", "suggested_category"]:
                    if key in changes:
                        entities["primary_pattern"][key] = changes[key]

            elif entity_type == "sub_pattern":
                for sp in entities.get("sub_patterns", []):
                    if sp.get("name") == entity_name:
                        if changes.get("fields"):
                            sp["fields"].update(changes["fields"])
                        for key in ["suggested_type", "confidence", "name", "suggested_category"]:
                            if key in changes:
                                sp[key] = changes[key]
                        break

            elif entity_type == "technology":
                for tech in entities.get("technologies", []):
                    if tech.get("name") == entity_name:
                        tech.update(changes)
                        break

            elif entity_type == "business_capability":
                for cap in entities.get("business_capabilities", []):
                    if cap.get("name") == entity_name:
                        cap.update(changes)
                        break

        elif action == "add":
            if entity_type == "sub_pattern":
                entities.setdefault("sub_patterns", []).append(changes)
            elif entity_type == "technology":
                entities.setdefault("technologies", []).append(changes)
            elif entity_type == "business_capability":
                entities.setdefault("business_capabilities", []).append(changes)

        elif action == "remove":
            if entity_type == "sub_pattern":
                entities["sub_patterns"] = [
                    sp for sp in entities.get("sub_patterns", [])
                    if sp.get("name") != entity_name
                ]
            elif entity_type == "technology":
                entities["technologies"] = [
                    t for t in entities.get("technologies", [])
                    if t.get("name") != entity_name
                ]
            elif entity_type == "business_capability":
                entities["business_capabilities"] = [
                    c for c in entities.get("business_capabilities", [])
                    if c.get("name") != entity_name
                ]

    return entities
