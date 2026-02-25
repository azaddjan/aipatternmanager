"""
DOCX Export Service.
Generates a Word document with cover page, TOC, and all patterns grouped by category.
"""
import io
import re
from datetime import datetime
from collections import defaultdict

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from services.neo4j_service import Neo4jService, BUILTIN_CATEGORIES


CATEGORY_ORDER = ["blueprint", "core", "intg", "agt", "kr", "xcut", "pip"]

CATEGORY_DESCRIPTIONS = {
    "blueprint": "Foundational structural patterns that define the overarching platform topology.",
    "core": "Core AI/LLM interaction patterns covering prompt engineering and model invocation.",
    "intg": "Integration patterns bridging the platform to AI model providers and tool ecosystems.",
    "agt": "Agent patterns combining reasoning, tool use, and memory for multi-step tasks.",
    "kr": "Knowledge & retrieval patterns covering the full RAG pipeline.",
    "xcut": "Cross-cutting concerns spanning governance, guardrails, and observability.",
    "pip": "Platform integration patterns for vendor portability and service contracts.",
}


class DocxExportService:
    """Builds a Word document export of the full pattern library."""

    def __init__(self, db: Neo4jService):
        self.db = db

    def generate_docx(self, team_ids=None, team_names=None) -> bytes:
        """Generate a DOCX file and return raw bytes."""
        self._team_names = team_names or []
        data = self._fetch_all_data(team_ids=team_ids)
        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # Heading styles
        for level in range(1, 4):
            h_style = doc.styles[f'Heading {level}']
            h_style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
            h_style.font.name = 'Calibri'

        self._add_cover_page(doc, data)
        self._add_toc(doc)
        self._set_update_fields_on_open(doc)
        self._add_page_numbers(doc)

        # Group patterns by category
        grouped = defaultdict(list)
        for p in data["patterns"]:
            grouped[p.get("category", "other")].append(p)

        cat_map = {c["code"]: c["label"] for c in data["categories"]}

        # Ordered categories
        seen = set()
        ordered_cats = []
        for c in CATEGORY_ORDER:
            if c in grouped:
                ordered_cats.append(c)
                seen.add(c)
        for c in sorted(grouped.keys()):
            if c not in seen:
                ordered_cats.append(c)

        # Pattern sections by category
        for cat_code in ordered_cats:
            cat_label = cat_map.get(cat_code, BUILTIN_CATEGORIES.get(cat_code, cat_code))
            pats = sorted(grouped[cat_code], key=lambda p: p["id"])

            # Category heading
            doc.add_page_break()
            doc.add_heading(cat_label, level=1)

            # Category description
            desc = CATEGORY_DESCRIPTIONS.get(cat_code, "")
            if desc:
                p = doc.add_paragraph(desc)
                p.style.font.size = Pt(10)
                run = p.runs[0] if p.runs else None
                if run:
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                    run.font.size = Pt(10)

            # Sort: AB first, then ABB, then SBB
            type_order = {"AB": 0, "ABB": 1, "SBB": 2}
            pats.sort(key=lambda p: (type_order.get(p.get("type", ""), 9), p["id"]))

            for pattern in pats:
                self._render_pattern(doc, pattern)

        # Technologies section
        if data["technologies"]:
            doc.add_page_break()
            doc.add_heading("Technology Registry", level=1)
            for tech in sorted(data["technologies"], key=lambda t: t.get("name", "")):
                self._render_technology(doc, tech)

        # PBC section
        if data["pbcs"]:
            doc.add_page_break()
            doc.add_heading("Business Capabilities (PBCs)", level=1)
            for pbc in sorted(data["pbcs"], key=lambda p: p.get("id", "")):
                self._render_pbc(doc, pbc)

        # Save to bytes
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ------------------------------------------------------------------
    # Data fetching
    # ------------------------------------------------------------------

    def _fetch_all_data(self, team_ids=None) -> dict:
        patterns, _ = self.db.list_patterns(limit=500, team_ids=team_ids)
        full_patterns = []
        for p in patterns:
            full = self.db.get_pattern_with_relationships(p["id"])
            if full:
                full_patterns.append(full)

        technologies, _ = self.db.list_technologies(limit=500)
        full_techs = []
        for t in technologies:
            full_t = self.db.get_technology_with_patterns(t["id"])
            if full_t:
                full_techs.append(full_t)

        pbcs = self.db.list_pbcs()
        categories = self.db.list_categories()

        return {
            "patterns": full_patterns,
            "technologies": full_techs,
            "pbcs": pbcs,
            "categories": categories,
        }

    # ------------------------------------------------------------------
    # Cover page
    # ------------------------------------------------------------------

    def _add_cover_page(self, doc, data):
        """Add a styled cover page."""
        # Spacer
        for _ in range(6):
            doc.add_paragraph("")

        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("Architecture Patterns")
        run.bold = True
        run.font.size = Pt(36)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("Pattern Catalog")
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x58, 0xa6, 0xff)

        doc.add_paragraph("")

        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        from zoneinfo import ZoneInfo
        est_now = datetime.now(ZoneInfo("America/New_York"))
        run = date_para.add_run(est_now.strftime("%B %d, %Y  %I:%M %p EST"))
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Team scope
        if self._team_names:
            scope_para = doc.add_paragraph()
            scope_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = scope_para.add_run("Teams: " + ", ".join(self._team_names))
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x58, 0xa6, 0xff)

        doc.add_paragraph("")

        # Stats summary
        patterns = data["patterns"]
        ab_count = sum(1 for p in patterns if p.get("type") == "AB")
        abb_count = sum(1 for p in patterns if p.get("type") == "ABB")
        sbb_count = sum(1 for p in patterns if p.get("type") == "SBB")
        tech_count = len(data["technologies"])
        pbc_count = len(data["pbcs"])

        stats = doc.add_paragraph()
        stats.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = stats.add_run(
            f"{ab_count} Blueprints  |  {abb_count} ABBs  |  {sbb_count} SBBs  |  "
            f"{tech_count} Technologies  |  {pbc_count} PBCs"
        )
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        doc.add_page_break()

    # ------------------------------------------------------------------
    # TOC and page numbers
    # ------------------------------------------------------------------

    def _add_toc(self, doc):
        """Insert a Word-native Table of Contents."""
        doc.add_heading("Table of Contents", level=1)
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        fldChar1 = parse_xml(r'<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls('w')))
        run._r.append(fldChar1)

        run2 = paragraph.add_run()
        instrText = parse_xml(
            r'<w:instrText {} xml:space="preserve"> TOC \o "1-2" \h \z \u </w:instrText>'.format(nsdecls('w'))
        )
        run2._r.append(instrText)

        run3 = paragraph.add_run()
        fldChar2 = parse_xml(r'<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls('w')))
        run3._r.append(fldChar2)

        doc.add_page_break()

    def _set_update_fields_on_open(self, doc):
        """Set the document to update fields (including TOC) on open."""
        settings = doc.settings.element
        update_fields = parse_xml(
            '<w:updateFields {} w:val="true"/>'.format(nsdecls('w'))
        )
        settings.append(update_fields)

    def _add_page_numbers(self, doc):
        """Add centered page numbers to the footer."""
        section = doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run()
        fldChar1 = parse_xml(r'<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls('w')))
        run._r.append(fldChar1)

        run2 = paragraph.add_run()
        instrText = parse_xml(
            r'<w:instrText {} xml:space="preserve"> PAGE </w:instrText>'.format(nsdecls('w'))
        )
        run2._r.append(instrText)

        run3 = paragraph.add_run()
        fldChar2 = parse_xml(r'<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls('w')))
        run3._r.append(fldChar2)

    # ------------------------------------------------------------------
    # Pattern rendering
    # ------------------------------------------------------------------

    def _render_pattern(self, doc, pattern):
        """Render a single pattern into the document."""
        pid = pattern["id"]
        ptype = pattern.get("type", "")
        name = pattern.get("name", "")

        # Pattern heading
        doc.add_heading(f"{pid} — {name}", level=2)

        # Metadata table
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Light Shading Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        team_name = pattern.get("team_name")
        meta_fields = [
            ("Type", ptype),
            ("Category", pattern.get("category", "")),
            ("Status", pattern.get("status", "")),
            ("Version", pattern.get("version", "")),
        ]
        if team_name:
            meta_fields.append(("Team", team_name))

        for label, value in meta_fields:
            if value:
                row = table.add_row()
                row.cells[0].text = label
                row.cells[1].text = str(value)
                # Bold the label
                for paragraph in row.cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

        doc.add_paragraph("")  # spacer

        # Tags
        tags = pattern.get("tags", [])
        if tags:
            para = doc.add_paragraph()
            run = para.add_run("Tags: ")
            run.bold = True
            run.font.size = Pt(9)
            run = para.add_run(", ".join(tags))
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x82, 0xF6)

        # Description (shared across all types)
        desc = pattern.get("description")
        if desc:
            doc.add_heading("Description", level=3)
            self._add_structured_text(doc, desc)

        # Deprecation note
        dep_note = pattern.get("deprecation_note")
        if dep_note:
            para = doc.add_paragraph()
            run = para.add_run("Deprecation Note: ")
            run.bold = True
            run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
            run = para.add_run(dep_note)
            run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

        # Type-specific content
        if ptype == "AB":
            self._render_ab_fields(doc, pattern)
        elif ptype == "ABB":
            self._render_abb_fields(doc, pattern)
        elif ptype == "SBB":
            self._render_sbb_fields(doc, pattern)

        # Diagrams (shared across all types)
        diagrams = pattern.get("diagrams", [])
        if diagrams:
            doc.add_heading("Diagrams", level=3)
            for diag in diagrams:
                title = diag.get("title", "Untitled")
                content = diag.get("content", "")
                para = doc.add_paragraph()
                run = para.add_run(title)
                run.bold = True
                if content:
                    code_para = doc.add_paragraph()
                    run = code_para.add_run(content)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Restrictions (shared across all types)
        restrictions = pattern.get("restrictions")
        if restrictions:
            para = doc.add_paragraph()
            run = para.add_run("Restrictions: ")
            run.bold = True
            run.font.color.rgb = RGBColor(0xF5, 0x9E, 0x0B)
            para.add_run(restrictions)

        # Relationships
        rels = pattern.get("relationships", [])
        if rels:
            doc.add_heading("Relationships", level=3)
            rel_table = doc.add_table(rows=1, cols=3)
            rel_table.style = 'Light Shading Accent 1'
            hdr = rel_table.rows[0].cells
            hdr[0].text = "Type"
            hdr[1].text = "Target"
            hdr[2].text = "Name"
            for r in rels:
                row = rel_table.add_row()
                row.cells[0].text = r.get("type") or ""
                row.cells[1].text = r.get("target_id") or ""
                row.cells[2].text = r.get("target_name") or ""

    def _render_ab_fields(self, doc, p):
        """Render Architecture Blueprint specific fields."""
        sections = [
            ("Intent", p.get("intent")),
            ("Problem", p.get("problem")),
            ("Solution", p.get("solution")),
            ("Structural Elements", p.get("structural_elements")),
            ("Invariants", p.get("invariants")),
            ("Inter-Element Contracts", p.get("inter_element_contracts")),
            ("Related Patterns", p.get("related_patterns_text")),
            ("Related ADRs", p.get("related_adrs")),
            ("Building Blocks Note", p.get("building_blocks_note")),
        ]
        for title, content in sections:
            if content:
                doc.add_heading(title, level=3)
                self._add_structured_text(doc, content)

    def _render_abb_fields(self, doc, p):
        """Render ABB specific fields."""
        if p.get("functionality"):
            doc.add_heading("Functionality", level=3)
            self._add_structured_text(doc, p["functionality"])

        if p.get("inbound_interfaces"):
            doc.add_heading("Inbound Interfaces", level=3)
            self._add_structured_text(doc, p["inbound_interfaces"])

        if p.get("outbound_interfaces"):
            doc.add_heading("Outbound Interfaces", level=3)
            self._add_structured_text(doc, p["outbound_interfaces"])

        if p.get("quality_attributes"):
            doc.add_heading("Quality Attributes", level=3)
            self._add_structured_text(doc, p["quality_attributes"])

        if p.get("compliance_requirements"):
            doc.add_heading("Compliance Requirements", level=3)
            self._add_structured_text(doc, p["compliance_requirements"])

        caps = p.get("business_capabilities", [])
        if caps:
            doc.add_heading("Business Capabilities", level=3)
            doc.add_paragraph(", ".join(caps))

    def _render_sbb_fields(self, doc, p):
        """Render SBB specific fields."""
        if p.get("specific_functionality"):
            doc.add_heading("Specific Functionality", level=3)
            self._add_structured_text(doc, p["specific_functionality"])

        # Solution Details table
        sol_fields = [
            ("Vendor", p.get("vendor")),
            ("Deployment Model", p.get("deployment_model")),
            ("Cost Tier", p.get("cost_tier")),
            ("Licensing", p.get("licensing")),
            ("Maturity", p.get("maturity")),
        ]
        sol_rows = [(label, val) for label, val in sol_fields if val]
        if sol_rows:
            doc.add_heading("Solution Details", level=3)
            table = doc.add_table(rows=0, cols=2)
            table.style = 'Light Shading Accent 1'
            for label, val in sol_rows:
                row = table.add_row()
                row.cells[0].text = label
                row.cells[1].text = str(val)
                for paragraph in row.cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            doc.add_paragraph("")  # spacer

        if p.get("inbound_interfaces"):
            doc.add_heading("Inbound Interfaces", level=3)
            self._add_structured_text(doc, p["inbound_interfaces"])

        if p.get("outbound_interfaces"):
            doc.add_heading("Outbound Interfaces", level=3)
            self._add_structured_text(doc, p["outbound_interfaces"])

        mapping = p.get("sbb_mapping", [])
        if mapping and isinstance(mapping, list):
            doc.add_heading("SBB Mapping", level=3)
            for row in mapping:
                if isinstance(row, dict):
                    key = row.get("key", "")
                    val = row.get("value", "")
                    para = doc.add_paragraph()
                    run = para.add_run(f"{key}: ")
                    run.bold = True
                    para.add_run(val)

    # ------------------------------------------------------------------
    # Technology rendering
    # ------------------------------------------------------------------

    def _render_technology(self, doc, tech):
        """Render a technology entry."""
        name = tech.get("name", tech["id"])
        doc.add_heading(name, level=2)

        table = doc.add_table(rows=0, cols=2)
        table.style = 'Light Shading Accent 1'

        fields = [
            ("ID", tech["id"]),
            ("Vendor", tech.get("vendor", "")),
            ("Category", tech.get("category", "")),
            ("Status", tech.get("status", "")),
            ("Cost Tier", tech.get("cost_tier", "")),
        ]
        for label, value in fields:
            if value:
                row = table.add_row()
                row.cells[0].text = label
                row.cells[1].text = str(value)
                for paragraph in row.cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

        if tech.get("description"):
            doc.add_paragraph("")
            doc.add_heading("Description", level=3)
            self._add_structured_text(doc, tech["description"])

        used_by = tech.get("used_by_patterns", [])
        if used_by:
            doc.add_heading("Used By Patterns", level=3)
            for p in used_by:
                if isinstance(p, dict):
                    doc.add_paragraph(f"{p.get('id', '')} — {p.get('name', '')}", style='List Bullet')
                elif isinstance(p, str):
                    doc.add_paragraph(p, style='List Bullet')

    # ------------------------------------------------------------------
    # PBC rendering
    # ------------------------------------------------------------------

    def _render_pbc(self, doc, pbc):
        """Render a PBC entry."""
        pid = pbc["id"]
        name = pbc.get("name", pid)
        doc.add_heading(f"{pid} — {name}", level=2)

        if pbc.get("description"):
            self._add_structured_text(doc, pbc["description"])

        abb_ids = pbc.get("abb_ids", [])
        if abb_ids:
            doc.add_heading("Composed ABBs", level=3)
            for abb_id in abb_ids:
                doc.add_paragraph(abb_id, style='List Bullet')

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def _add_structured_text(self, doc, text):
        """Add text to the document, handling lists and tables from markdown-style content."""
        if not text:
            return

        lines = text.split("\n")
        in_table = False
        table_rows = []

        for line in lines:
            stripped = line.strip()
            is_table_row = stripped.startswith("|") and stripped.endswith("|") and len(stripped) > 2

            if is_table_row:
                # Skip separator rows
                cleaned = stripped.replace("|", "").replace("-", "").replace(":", "").replace(" ", "")
                if cleaned:
                    table_rows.append(stripped)
                in_table = True
            else:
                if in_table and table_rows:
                    self._flush_table_to_doc(doc, table_rows)
                    table_rows = []
                    in_table = False

                if stripped.startswith("- ") or stripped.startswith("* "):
                    item_text = stripped[2:]
                    # Clean bold markers
                    item_text = re.sub(r'\*\*(.+?)\*\*', r'\1', item_text)
                    doc.add_paragraph(item_text, style='List Bullet')
                elif stripped:
                    # Clean bold markers for paragraphs
                    clean_text = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped)
                    doc.add_paragraph(clean_text)

        # Flush remaining table
        if table_rows:
            self._flush_table_to_doc(doc, table_rows)

    def _flush_table_to_doc(self, doc, rows):
        """Convert markdown table rows into a Word table."""
        if not rows:
            return

        def parse_cells(row_str):
            cells = row_str.split("|")
            if cells and not cells[0].strip():
                cells = cells[1:]
            if cells and not cells[-1].strip():
                cells = cells[:-1]
            return [c.strip() for c in cells]

        header_cells = parse_cells(rows[0])
        col_count = len(header_cells)
        if col_count == 0:
            return

        table = doc.add_table(rows=1, cols=col_count)
        table.style = 'Light Shading Accent 1'

        # Header row
        hdr = table.rows[0].cells
        for i, cell in enumerate(header_cells):
            if i < col_count:
                clean = re.sub(r'\*\*(.+?)\*\*', r'\1', cell)
                hdr[i].text = clean
                for paragraph in hdr[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

        # Data rows
        for row_str in rows[1:]:
            row_cells = parse_cells(row_str)
            row = table.add_row()
            for i, cell in enumerate(row_cells):
                if i < col_count:
                    clean = re.sub(r'\*\*(.+?)\*\*', r'\1', cell)
                    row.cells[i].text = clean

        doc.add_paragraph("")  # spacer after table
