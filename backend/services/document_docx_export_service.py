"""
Document DOCX Export Service.
Generates a Word document from a single Document with cover page, TOC, sections, and linked entities.
"""
import io
import re
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


DOC_TYPE_LABELS = {
    "guide": "Architecture Guide",
    "reference": "Reference Document",
    "adr": "Architecture Decision Record",
    "overview": "Architecture Overview",
    "other": "Document",
}


class DocumentDocxExportService:
    """Builds a Word document export of a single Document."""

    def generate_docx(self, doc: dict) -> bytes:
        """Generate a DOCX file from a document dict (as returned by get_document)."""
        doc_obj = Document()

        # Set default font
        style = doc_obj.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # Heading styles
        for level in range(1, 4):
            h_style = doc_obj.styles[f"Heading {level}"]
            h_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            h_style.font.name = "Calibri"

        self._add_cover_page(doc_obj, doc)
        self._add_toc(doc_obj)
        self._set_update_fields_on_open(doc_obj)
        self._add_page_numbers(doc_obj)

        # Render sections
        sections = doc.get("sections", [])
        for i, section in enumerate(sections):
            if i > 0:
                doc_obj.add_page_break()
            self._render_section(doc_obj, section)

        # Linked entities appendix
        linked = doc.get("linked_entities", [])
        if linked:
            self._render_linked_entities(doc_obj, linked)

        # Save to bytes
        buf = io.BytesIO()
        doc_obj.save(buf)
        return buf.getvalue()

    # ------------------------------------------------------------------
    # Cover page
    # ------------------------------------------------------------------

    def _add_cover_page(self, doc_obj, doc: dict):
        """Add a styled cover page."""
        # Spacer
        for _ in range(6):
            doc_obj.add_paragraph("")

        # Title
        title_para = doc_obj.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(doc.get("title", "Untitled Document"))
        run.bold = True
        run.font.size = Pt(36)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        # Doc type subtitle
        doc_type = doc.get("doc_type", "other")
        type_label = DOC_TYPE_LABELS.get(doc_type, doc_type.title())
        subtitle = doc_obj.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(type_label)
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x58, 0xA6, 0xFF)

        # Target audience
        audience = doc.get("target_audience", "")
        if audience:
            audience_para = doc_obj.add_paragraph()
            audience_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = audience_para.add_run(f"For {audience}")
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc_obj.add_paragraph("")

        # Date and author
        date_para = doc_obj.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        from zoneinfo import ZoneInfo
        est_now = datetime.now(ZoneInfo("America/New_York"))
        date_text = est_now.strftime("%B %d, %Y  %I:%M %p EST")
        created_by = doc.get("created_by", "")
        if created_by:
            date_text += f"  |  {created_by}"
        run = date_para.add_run(date_text)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Team
        team_name = doc.get("team_name")
        if team_name:
            team_para = doc_obj.add_paragraph()
            team_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = team_para.add_run(f"Team: {team_name}")
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x58, 0xA6, 0xFF)

        # Summary
        summary = doc.get("summary", "")
        if summary:
            doc_obj.add_paragraph("")
            summary_para = doc_obj.add_paragraph()
            summary_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = summary_para.add_run(summary)
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        doc_obj.add_paragraph("")

        # Stats
        section_count = len(doc.get("sections", []))
        link_count = len(doc.get("linked_entities", []))
        stats_para = doc_obj.add_paragraph()
        stats_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        stats_text = f"{section_count} section{'s' if section_count != 1 else ''}"
        stats_text += f"  |  {link_count} linked entit{'ies' if link_count != 1 else 'y'}"
        # Tags
        tags = doc.get("tags", [])
        if tags:
            stats_text += f"  |  Tags: {', '.join(tags)}"
        run = stats_para.add_run(stats_text)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # Status badge
        status = doc.get("status", "draft")
        status_para = doc_obj.add_paragraph()
        status_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = status_para.add_run(f"Status: {status.upper()}")
        run.font.size = Pt(10)
        run.bold = True
        if status == "published":
            run.font.color.rgb = RGBColor(0x22, 0xC5, 0x5E)
        elif status == "archived":
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        else:
            run.font.color.rgb = RGBColor(0xEA, 0xB3, 0x08)

        doc_obj.add_page_break()

    # ------------------------------------------------------------------
    # TOC and page numbers
    # ------------------------------------------------------------------

    def _add_toc(self, doc_obj):
        """Insert a Word-native Table of Contents."""
        doc_obj.add_heading("Table of Contents", level=1)
        paragraph = doc_obj.add_paragraph()
        run = paragraph.add_run()
        fldChar1 = parse_xml(
            r'<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls("w"))
        )
        run._r.append(fldChar1)

        run2 = paragraph.add_run()
        instrText = parse_xml(
            r'<w:instrText {} xml:space="preserve"> TOC \o "1-2" \h \z \u </w:instrText>'.format(
                nsdecls("w")
            )
        )
        run2._r.append(instrText)

        run3 = paragraph.add_run()
        fldChar2 = parse_xml(
            r'<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls("w"))
        )
        run3._r.append(fldChar2)

        doc_obj.add_page_break()

    def _set_update_fields_on_open(self, doc_obj):
        """Set the document to update fields (including TOC) on open."""
        settings = doc_obj.settings.element
        update_fields = parse_xml(
            '<w:updateFields {} w:val="true"/>'.format(nsdecls("w"))
        )
        settings.append(update_fields)

    def _add_page_numbers(self, doc_obj):
        """Add centered page numbers to the footer."""
        section = doc_obj.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run()
        fldChar1 = parse_xml(
            r'<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls("w"))
        )
        run._r.append(fldChar1)

        run2 = paragraph.add_run()
        instrText = parse_xml(
            r'<w:instrText {} xml:space="preserve"> PAGE </w:instrText>'.format(
                nsdecls("w")
            )
        )
        run2._r.append(instrText)

        run3 = paragraph.add_run()
        fldChar2 = parse_xml(
            r'<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls("w"))
        )
        run3._r.append(fldChar2)

    # ------------------------------------------------------------------
    # Section rendering
    # ------------------------------------------------------------------

    def _render_section(self, doc_obj, section: dict):
        """Render a single document section with its content."""
        title = section.get("title", "Untitled Section")
        content = section.get("content", "")

        doc_obj.add_heading(title, level=1)

        if content:
            self._add_structured_text(doc_obj, content)
        else:
            p = doc_obj.add_paragraph("(No content)")
            run = p.runs[0]
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run.italic = True

    # ------------------------------------------------------------------
    # Linked entities appendix
    # ------------------------------------------------------------------

    def _render_linked_entities(self, doc_obj, entities: list):
        """Render a linked entities appendix as a table."""
        doc_obj.add_page_break()
        doc_obj.add_heading("Referenced Entities", level=1)

        intro = doc_obj.add_paragraph(
            "The following catalog entities are referenced by this document."
        )
        run = intro.runs[0] if intro.runs else None
        if run:
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            run.font.size = Pt(10)

        doc_obj.add_paragraph("")

        # Sort: Patterns first, then Technologies, then PBCs, then others
        label_order = {"Pattern": 0, "Technology": 1, "PBC": 2}
        sorted_entities = sorted(
            entities,
            key=lambda e: (label_order.get(e.get("label", ""), 9), e.get("id", "")),
        )

        # Create table
        table = doc_obj.add_table(rows=1, cols=3)
        table.style = "Light Shading Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Header row
        hdr = table.rows[0].cells
        hdr[0].text = "Type"
        hdr[1].text = "ID"
        hdr[2].text = "Name"
        for cell in hdr:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Data rows
        for entity in sorted_entities:
            row = table.add_row()
            row.cells[0].text = entity.get("label", "")
            row.cells[1].text = entity.get("id", "")
            row.cells[2].text = entity.get("name", "")

    # ------------------------------------------------------------------
    # Structured text rendering (markdown-aware)
    # ------------------------------------------------------------------

    def _add_structured_text(self, doc_obj, text: str):
        """Add text to the document, handling markdown lists, tables, code blocks, and headings."""
        if not text:
            return

        lines = text.split("\n")
        in_table = False
        table_rows = []
        in_code_block = False
        code_language = ""
        code_lines = []

        for line in lines:
            stripped = line.strip()

            # Code block detection (``` or ```language)
            if stripped.startswith("```"):
                if not in_code_block:
                    # Starting a code block
                    in_code_block = True
                    code_language = stripped[3:].strip()
                    code_lines = []

                    # Flush any pending table
                    if in_table and table_rows:
                        self._flush_table_to_doc(doc_obj, table_rows)
                        table_rows = []
                        in_table = False
                else:
                    # Ending a code block
                    in_code_block = False
                    self._render_code_block(doc_obj, code_lines, code_language)
                    code_lines = []
                    code_language = ""
                continue

            if in_code_block:
                code_lines.append(line)
                continue

            # Table row detection
            is_table_row = (
                stripped.startswith("|")
                and stripped.endswith("|")
                and len(stripped) > 2
            )

            if is_table_row:
                # Skip separator rows (e.g. |---|---|)
                cleaned = (
                    stripped.replace("|", "")
                    .replace("-", "")
                    .replace(":", "")
                    .replace(" ", "")
                )
                if cleaned:
                    table_rows.append(stripped)
                in_table = True
            else:
                # Flush table if we were in one
                if in_table and table_rows:
                    self._flush_table_to_doc(doc_obj, table_rows)
                    table_rows = []
                    in_table = False

                # Markdown headings (### Heading)
                if stripped.startswith("### "):
                    doc_obj.add_heading(stripped[4:], level=3)
                elif stripped.startswith("## "):
                    doc_obj.add_heading(stripped[3:], level=2)
                elif stripped.startswith("# "):
                    doc_obj.add_heading(stripped[2:], level=2)
                # Bullet lists
                elif stripped.startswith("- ") or stripped.startswith("* "):
                    item_text = stripped[2:]
                    item_text = self._clean_markdown_inline(item_text)
                    doc_obj.add_paragraph(item_text, style="List Bullet")
                # Numbered lists
                elif re.match(r"^\d+\.\s", stripped):
                    item_text = re.sub(r"^\d+\.\s", "", stripped)
                    item_text = self._clean_markdown_inline(item_text)
                    doc_obj.add_paragraph(item_text, style="List Number")
                # Horizontal rule
                elif stripped in ("---", "***", "___"):
                    # Add a thin horizontal line
                    doc_obj.add_paragraph("_" * 60)
                # Regular paragraph
                elif stripped:
                    clean_text = self._clean_markdown_inline(stripped)
                    para = doc_obj.add_paragraph()
                    self._add_rich_text(para, clean_text)

        # Flush remaining table
        if table_rows:
            self._flush_table_to_doc(doc_obj, table_rows)

        # Flush remaining code block (unclosed)
        if in_code_block and code_lines:
            self._render_code_block(doc_obj, code_lines, code_language)

    def _render_code_block(self, doc_obj, lines: list, language: str = ""):
        """Render a code block as monospace text with a label."""
        if language.lower() == "mermaid":
            label_text = "Mermaid Diagram"
        elif language:
            label_text = f"Code ({language})"
        else:
            label_text = "Code"

        # Label
        label_para = doc_obj.add_paragraph()
        run = label_para.add_run(label_text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x58, 0xA6, 0xFF)

        # Code content
        code_text = "\n".join(lines)
        if code_text.strip():
            code_para = doc_obj.add_paragraph()
            run = code_para.add_run(code_text)
            run.font.name = "Consolas"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        doc_obj.add_paragraph("")  # spacer

    def _flush_table_to_doc(self, doc_obj, rows: list):
        """Convert markdown table rows into a Word table."""
        if not rows:
            return

        def parse_cells(row_str):
            cells = row_str.split("|")
            if cells and not cells[0].strip():
                cells = cells[1:]
            if cells and not cells[-1].strip():
                cells = cells[:-1]
            return [c.strip() for c in cells]

        header_cells = parse_cells(rows[0])
        col_count = len(header_cells)
        if col_count == 0:
            return

        table = doc_obj.add_table(rows=1, cols=col_count)
        table.style = "Light Shading Accent 1"

        # Header row
        hdr = table.rows[0].cells
        for i, cell in enumerate(header_cells):
            if i < col_count:
                clean = self._clean_markdown_inline(cell)
                hdr[i].text = clean
                for paragraph in hdr[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

        # Data rows
        for row_str in rows[1:]:
            row_cells = parse_cells(row_str)
            row = table.add_row()
            for i, cell in enumerate(row_cells):
                if i < col_count:
                    clean = self._clean_markdown_inline(cell)
                    row.cells[i].text = clean

        doc_obj.add_paragraph("")  # spacer after table

    def _clean_markdown_inline(self, text: str) -> str:
        """Remove inline markdown formatting for plain text."""
        # Bold
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        # Italic
        text = re.sub(r"\*(.+?)\*", r"\1", text)
        text = re.sub(r"_(.+?)_", r"\1", text)
        # Inline code
        text = re.sub(r"`(.+?)`", r"\1", text)
        # Links [text](url) -> text (url)
        text = re.sub(r"\[(.+?)\]\((.+?)\)", r"\1 (\2)", text)
        return text

    def _add_rich_text(self, paragraph, text: str):
        """Add text with basic bold/italic rendering to a paragraph."""
        # For simplicity, add as a single run with cleaned text
        # A more advanced version could parse **bold** and _italic_ into separate runs
        run = paragraph.add_run(text)
        run.font.size = Pt(11)
