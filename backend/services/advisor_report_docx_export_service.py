"""Advisor Report DOCX Export Service — generates a Word document for a single report."""
import io
import re
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


class AdvisorReportDocxExportService:
    """Generates a Word document for a single advisor report."""

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def generate_docx(self, report: dict) -> bytes:
        """Generate a DOCX file for a single advisor report and return raw bytes."""
        result = report.get("result_json", {})
        analysis = result.get("analysis", {})

        doc = Document()
        self._set_styles(doc)
        self._add_cover_page(doc, report, analysis)
        self._add_summary(doc, analysis)
        self._add_recommended_patterns(doc, analysis)
        self._add_comparisons(doc, analysis)
        self._add_architecture(doc, analysis)
        self._add_gaps(doc, analysis)
        self._add_vector_matches(doc, result)
        self._add_page_numbers(doc)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ------------------------------------------------------------------
    # Styles
    # ------------------------------------------------------------------

    def _set_styles(self, doc: Document) -> None:
        """Set Normal style to Calibri 11pt, heading styles with brand colours."""
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        for level in range(1, 4):
            h_style = doc.styles[f"Heading {level}"]
            h_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            h_style.font.name = "Calibri"

    # ------------------------------------------------------------------
    # Cover page
    # ------------------------------------------------------------------

    def _add_cover_page(self, doc: Document, report: dict, analysis: dict) -> None:
        """Add a styled cover page with report metadata."""
        # Spacer
        for _ in range(6):
            doc.add_paragraph("")

        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run("Advisor Report")
        run.bold = True
        run.font.size = Pt(36)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        # Subtitle — report ID and confidence
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        report_id = report.get("id", "")
        confidence = report.get("confidence", analysis.get("confidence", ""))
        subtitle_text = report_id
        if confidence:
            subtitle_text += f"  |  Confidence: {confidence}"
        run = subtitle_para.add_run(subtitle_text)
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x58, 0xA6, 0xFF)

        # Report title
        if report.get("title"):
            title_name = doc.add_paragraph()
            title_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title_name.add_run(report["title"])
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        doc.add_paragraph("")

        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        created_at = report.get("created_at", "")
        if created_at:
            try:
                dt = datetime.fromisoformat(created_at)
                date_str = dt.strftime("%B %d, %Y")
            except (ValueError, TypeError):
                date_str = datetime.now().strftime("%B %d, %Y")
        else:
            date_str = datetime.now().strftime("%B %d, %Y")
        run = date_para.add_run(date_str)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph("")

        # Problem description
        problem = report.get("problem", "")
        if problem:
            problem_para = doc.add_paragraph()
            problem_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = problem_para.add_run("Problem Statement")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

            desc_para = doc.add_paragraph()
            desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = desc_para.add_run(problem)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        doc.add_paragraph("")

        # Provider / model info
        provider = report.get("provider", "")
        model = report.get("model", "")
        category_focus = report.get("category_focus", "")
        tech_prefs = report.get("technology_preferences", [])
        starred = report.get("starred", False)

        info_parts = []
        if provider:
            info_parts.append(f"Provider: {provider}")
        if model:
            info_parts.append(f"Model: {model}")
        if category_focus:
            info_parts.append(f"Category Focus: {category_focus}")
        if tech_prefs:
            info_parts.append(f"Tech Preferences: {', '.join(tech_prefs)}")
        if starred:
            info_parts.append("Starred")

        if info_parts:
            info_para = doc.add_paragraph()
            info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = info_para.add_run("  |  ".join(info_parts))
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        doc.add_page_break()

    # ------------------------------------------------------------------
    # Executive Summary
    # ------------------------------------------------------------------

    def _add_summary(self, doc: Document, analysis: dict) -> None:
        """Add the Executive Summary section."""
        summary_text = analysis.get("summary", "")
        if not summary_text:
            return

        doc.add_heading("Executive Summary", level=1)
        self._add_structured_text(doc, summary_text)

    # ------------------------------------------------------------------
    # Recommended Patterns
    # ------------------------------------------------------------------

    def _add_recommended_patterns(self, doc: Document, analysis: dict) -> None:
        """Add Recommended Patterns with PBC, ABB, and SBB sub-sections."""
        pbcs = analysis.get("recommended_pbcs", [])
        abbs = analysis.get("recommended_abbs", [])
        sbbs = analysis.get("recommended_sbbs", [])

        if not pbcs and not abbs and not sbbs:
            return

        doc.add_heading("Recommended Patterns", level=1)

        # --- PBCs ---
        if pbcs:
            doc.add_heading("Recommended PBCs", level=2)
            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Shading Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.LEFT

            hdr = table.rows[0].cells
            for i, label in enumerate(["ID", "Name", "Confidence", "Relevance"]):
                hdr[i].text = label
                for paragraph in hdr[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            for pbc in pbcs:
                row = table.add_row()
                row.cells[0].text = pbc.get("id", "")
                row.cells[1].text = pbc.get("name", "")
                row.cells[2].text = pbc.get("confidence", "")
                row.cells[3].text = pbc.get("relevance", "")

            doc.add_paragraph("")  # spacer

        # --- ABBs ---
        if abbs:
            doc.add_heading("Recommended ABBs", level=2)
            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Shading Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.LEFT

            hdr = table.rows[0].cells
            for i, label in enumerate(["ID", "Name", "Confidence", "Role"]):
                hdr[i].text = label
                for paragraph in hdr[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            for abb in abbs:
                row = table.add_row()
                row.cells[0].text = abb.get("id", "")
                row.cells[1].text = abb.get("name", "")
                row.cells[2].text = abb.get("confidence", "")
                row.cells[3].text = abb.get("role", "")

            doc.add_paragraph("")  # spacer

        # --- SBBs ---
        if sbbs:
            doc.add_heading("Recommended SBBs", level=2)
            table = doc.add_table(rows=1, cols=5)
            table.style = "Light Shading Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.LEFT

            hdr = table.rows[0].cells
            for i, label in enumerate(["ID", "Name", "Confidence", "Justification", "Technologies"]):
                hdr[i].text = label
                for paragraph in hdr[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            for sbb in sbbs:
                row = table.add_row()
                row.cells[0].text = sbb.get("id", "")
                row.cells[1].text = sbb.get("name", "")
                row.cells[2].text = sbb.get("confidence", "")
                justification = sbb.get("justification", "")
                restrictions = sbb.get("restrictions_note", "")
                if restrictions:
                    justification += f" [{restrictions}]"
                row.cells[3].text = justification
                techs = sbb.get("technologies", [])
                row.cells[4].text = ", ".join(techs) if isinstance(techs, list) else str(techs)

            doc.add_paragraph("")  # spacer

    # ------------------------------------------------------------------
    # SBB Comparisons
    # ------------------------------------------------------------------

    def _add_comparisons(self, doc: Document, analysis: dict) -> None:
        """Add the SBB Comparisons section."""
        comparisons = analysis.get("sbb_comparisons", [])
        if not comparisons:
            return

        doc.add_heading("SBB Comparisons", level=1)

        for group in comparisons:
            context = group.get("context", "Comparison")
            doc.add_heading(context, level=2)

            sbbs = group.get("sbbs", [])
            if sbbs:
                table = doc.add_table(rows=1, cols=4)
                table.style = "Light Shading Accent 1"
                table.alignment = WD_TABLE_ALIGNMENT.LEFT

                hdr = table.rows[0].cells
                for i, label in enumerate(["SBB", "Strengths", "Weaknesses", "Best For"]):
                    hdr[i].text = label
                    for paragraph in hdr[i].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

                for sbb in sbbs:
                    row = table.add_row()
                    sbb_label = sbb.get("id", "")
                    sbb_name = sbb.get("name", "")
                    if sbb_label and sbb_name:
                        row.cells[0].text = f"{sbb_label} — {sbb_name}"
                    else:
                        row.cells[0].text = sbb_label or sbb_name

                    strengths = sbb.get("strengths", [])
                    row.cells[1].text = "\n".join(strengths) if isinstance(strengths, list) else str(strengths)

                    weaknesses = sbb.get("weaknesses", [])
                    row.cells[2].text = "\n".join(weaknesses) if isinstance(weaknesses, list) else str(weaknesses)

                    row.cells[3].text = sbb.get("best_for", "")

                doc.add_paragraph("")  # spacer

            recommendation = group.get("recommendation", "")
            if recommendation:
                rec_para = doc.add_paragraph()
                run = rec_para.add_run("Recommendation: ")
                run.bold = True
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
                rec_para.add_run(recommendation)

    # ------------------------------------------------------------------
    # Architecture
    # ------------------------------------------------------------------

    def _add_architecture(self, doc: Document, analysis: dict) -> None:
        """Add Architecture Composition and Data Flow sections."""
        composition = analysis.get("architecture_composition", "")
        data_flow = analysis.get("data_flow", "")

        if not composition and not data_flow:
            return

        doc.add_heading("Architecture", level=1)

        if composition:
            doc.add_heading("Architecture Composition", level=2)
            self._add_structured_text(doc, composition)

        if data_flow:
            doc.add_heading("Data Flow", level=2)
            self._add_structured_text(doc, data_flow)

    # ------------------------------------------------------------------
    # Platform Gaps
    # ------------------------------------------------------------------

    def _add_gaps(self, doc: Document, analysis: dict) -> None:
        """Add Platform Gaps section."""
        gaps = analysis.get("platform_gaps", [])
        if not gaps:
            return

        doc.add_heading("Platform Gaps", level=1)

        for gap in gaps:
            capability = gap.get("capability", "Unknown Capability")
            doc.add_heading(capability, level=2)

            # Priority and category line
            priority = gap.get("priority", "")
            category = gap.get("category", "")
            meta_parts = []
            if priority:
                meta_parts.append(f"Priority: {priority}")
            if category:
                meta_parts.append(f"Category: {category}")
            if meta_parts:
                meta_para = doc.add_paragraph()
                run = meta_para.add_run("  |  ".join(meta_parts))
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x58, 0xA6, 0xFF)

            # Description
            description = gap.get("description", "")
            if description:
                doc.add_paragraph(description)

            # Rationale
            rationale = gap.get("rationale", "")
            if rationale:
                rationale_para = doc.add_paragraph()
                run = rationale_para.add_run("Rationale: ")
                run.bold = True
                rationale_para.add_run(rationale)

    # ------------------------------------------------------------------
    # Appendix — Vector Similarity Matches
    # ------------------------------------------------------------------

    def _add_vector_matches(self, doc: Document, result: dict) -> None:
        """Add Appendix with vector similarity match results."""
        vector_matches = result.get("vector_matches", {})
        patterns = vector_matches.get("patterns", [])
        technologies = vector_matches.get("technologies", [])
        pbcs = vector_matches.get("pbcs", [])

        all_rows = []
        for p in patterns:
            all_rows.append({
                "id": p.get("id", ""),
                "name": p.get("name", ""),
                "type": p.get("type", p.get("category", "")),
                "score": p.get("score", 0),
            })
        for t in technologies:
            all_rows.append({
                "id": t.get("id", ""),
                "name": t.get("name", ""),
                "type": f"Technology ({t.get('vendor', '')})" if t.get("vendor") else "Technology",
                "score": t.get("score", 0),
            })
        for pbc in pbcs:
            all_rows.append({
                "id": pbc.get("id", ""),
                "name": pbc.get("name", ""),
                "type": "PBC",
                "score": pbc.get("score", 0),
            })

        if not all_rows:
            return

        # Sort by score descending
        all_rows.sort(key=lambda r: r["score"], reverse=True)

        doc.add_page_break()
        doc.add_heading("Appendix: Vector Similarity Matches", level=1)

        # Graph stats summary
        graph_stats = result.get("graph_stats", {})
        if graph_stats:
            stats_parts = []
            for key in ["pbcs", "abbs", "sbbs", "technologies"]:
                count = graph_stats.get(key, 0)
                if count:
                    stats_parts.append(f"{key.upper()}: {count}")
            if stats_parts:
                stats_para = doc.add_paragraph()
                run = stats_para.add_run("Graph Stats: " + "  |  ".join(stats_parts))
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                doc.add_paragraph("")  # spacer

        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Shading Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        hdr = table.rows[0].cells
        for i, label in enumerate(["ID", "Name", "Type", "Score"]):
            hdr[i].text = label
            for paragraph in hdr[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for item in all_rows:
            row = table.add_row()
            row.cells[0].text = item["id"]
            row.cells[1].text = item["name"]
            row.cells[2].text = item["type"]
            score = item["score"]
            row.cells[3].text = f"{score:.2f}" if isinstance(score, (int, float)) else str(score)

        doc.add_paragraph("")  # spacer

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def _add_structured_text(self, doc: Document, text: str) -> None:
        """Add text to the document, handling markdown lists, bold, and tables."""
        if not text:
            return

        lines = text.split("\n")
        in_table = False
        table_rows = []

        for line in lines:
            stripped = line.strip()
            is_table_row = stripped.startswith("|") and stripped.endswith("|") and len(stripped) > 2

            if is_table_row:
                # Skip separator rows (lines like |---|---|)
                cleaned = stripped.replace("|", "").replace("-", "").replace(":", "").replace(" ", "")
                if cleaned:
                    table_rows.append(stripped)
                in_table = True
            else:
                if in_table and table_rows:
                    self._flush_table_to_doc(doc, table_rows)
                    table_rows = []
                    in_table = False

                if stripped.startswith("- ") or stripped.startswith("* "):
                    item_text = stripped[2:]
                    item_text = self._apply_bold_runs_to_paragraph(doc, item_text, style="List Bullet")
                elif stripped.startswith("### "):
                    doc.add_heading(stripped[4:], level=3)
                elif stripped.startswith("## "):
                    doc.add_heading(stripped[3:], level=2)
                elif stripped.startswith("# "):
                    doc.add_heading(stripped[2:], level=1)
                elif stripped:
                    self._apply_bold_runs_to_paragraph(doc, stripped)

        # Flush remaining table
        if table_rows:
            self._flush_table_to_doc(doc, table_rows)

    def _apply_bold_runs_to_paragraph(self, doc: Document, text: str, style: str = None) -> None:
        """Create a paragraph with bold runs for **text** segments."""
        if style:
            para = doc.add_paragraph(style=style)
        else:
            para = doc.add_paragraph()

        # Split on bold markers: **bold text**
        parts = re.split(r'(\*\*.+?\*\*)', text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = para.add_run(part[2:-2])
                run.bold = True
            else:
                para.add_run(part)

    def _flush_table_to_doc(self, doc: Document, rows: list) -> None:
        """Convert markdown table rows into a Word table."""
        if not rows:
            return

        def parse_cells(row_str: str) -> list:
            cells = row_str.split("|")
            if cells and not cells[0].strip():
                cells = cells[1:]
            if cells and not cells[-1].strip():
                cells = cells[:-1]
            return [c.strip() for c in cells]

        header_cells = parse_cells(rows[0])
        col_count = len(header_cells)
        if col_count == 0:
            return

        table = doc.add_table(rows=1, cols=col_count)
        table.style = "Light Shading Accent 1"

        # Header row
        hdr = table.rows[0].cells
        for i, cell in enumerate(header_cells):
            if i < col_count:
                clean = re.sub(r'\*\*(.+?)\*\*', r'\1', cell)
                hdr[i].text = clean
                for paragraph in hdr[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

        # Data rows
        for row_str in rows[1:]:
            row_cells = parse_cells(row_str)
            row = table.add_row()
            for i, cell in enumerate(row_cells):
                if i < col_count:
                    clean = re.sub(r'\*\*(.+?)\*\*', r'\1', cell)
                    row.cells[i].text = clean

        doc.add_paragraph("")  # spacer after table

    def _add_page_numbers(self, doc: Document) -> None:
        """Add centered page numbers to the footer."""
        section = doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run()
        fldChar1 = parse_xml(
            r'<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls('w'))
        )
        run._r.append(fldChar1)

        run2 = paragraph.add_run()
        instrText = parse_xml(
            r'<w:instrText {} xml:space="preserve"> PAGE </w:instrText>'.format(nsdecls('w'))
        )
        run2._r.append(instrText)

        run3 = paragraph.add_run()
        fldChar2 = parse_xml(
            r'<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls('w'))
        )
        run3._r.append(fldChar2)
