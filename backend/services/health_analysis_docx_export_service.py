"""Health Analysis DOCX Export Service — generates a Word document for a single health analysis."""
import io
import re
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# The 9 analysis areas produced by the AI deep analysis prompt
ANALYSIS_AREAS = [
    ("Architecture Coherence", "architecture_coherence"),
    ("ABB\u2013SBB Alignment", "abb_sbb_alignment"),
    ("Interface Consistency", "interface_consistency"),
    ("Business Capability Gaps", "business_capability_gaps"),
    ("Vendor & Technology Risk", "vendor_technology_risk"),
    ("Content Quality", "content_quality"),
    ("Cross-Pattern Overlap", "cross_pattern_overlap"),
    ("PBC Composition", "pbc_composition"),
]

# Keys that contain recommendation-like lists (excluded from "findings" aggregation)
_RECOMMENDATION_KEYS = {"recommendations", "consolidation_suggestions"}


class HealthAnalysisDocxExportService:
    """Generates a Word document for a single health analysis."""

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def generate_docx(self, analysis: dict) -> bytes:
        """Generate a DOCX file for a single health analysis and return raw bytes."""
        analysis_json = analysis.get("analysis_json", {})

        doc = Document()
        self._set_styles(doc)
        self._add_cover_page(doc, analysis)

        # If analysis_json is a dict with structured data, render sections
        if isinstance(analysis_json, dict) and "raw_text" not in analysis_json:
            self._add_executive_summary(doc, analysis, analysis_json)
            self._add_assessment_scores(doc, analysis_json)
            self._add_detailed_assessments(doc, analysis_json)
            self._add_maturity_roadmap(doc, analysis_json)
        else:
            # Fallback: raw_text key present or analysis_json is a string
            self._add_raw_text_fallback(doc, analysis_json)

        self._add_page_numbers(doc)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ------------------------------------------------------------------
    # Styles
    # ------------------------------------------------------------------

    def _set_styles(self, doc: Document) -> None:
        """Set Normal style to Calibri 11pt, heading styles with brand colours."""
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        for level in range(1, 4):
            h_style = doc.styles[f"Heading {level}"]
            h_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            h_style.font.name = "Calibri"

    # ------------------------------------------------------------------
    # Cover page
    # ------------------------------------------------------------------

    def _add_cover_page(self, doc: Document, analysis: dict) -> None:
        """Add a styled cover page with analysis metadata."""
        # Spacer
        for _ in range(6):
            doc.add_paragraph("")

        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run("Health Analysis")
        run.bold = True
        run.font.size = Pt(36)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        # Subtitle — analysis ID and health score
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        analysis_id = analysis.get("id", "")
        health_score = analysis.get("health_score")
        subtitle_text = analysis_id
        if health_score is not None:
            subtitle_text += f"  |  Health Score: {health_score}"
        run = subtitle_para.add_run(subtitle_text)
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x58, 0xA6, 0xFF)

        # Analysis title
        if analysis.get("title"):
            title_name = doc.add_paragraph()
            title_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title_name.add_run(analysis["title"])
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        doc.add_paragraph("")

        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        created_at = analysis.get("created_at", "")
        if created_at:
            try:
                dt = datetime.fromisoformat(created_at)
                date_str = dt.strftime("%B %d, %Y")
            except (ValueError, TypeError):
                date_str = datetime.now().strftime("%B %d, %Y")
        else:
            date_str = datetime.now().strftime("%B %d, %Y")
        run = date_para.add_run(date_str)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph("")

        # Provider / model / pattern count info
        provider = analysis.get("provider", "")
        model = analysis.get("model", "")
        pattern_count = analysis.get("pattern_count")

        info_parts = []
        if provider:
            info_parts.append(f"Provider: {provider}")
        if model:
            info_parts.append(f"Model: {model}")
        if pattern_count is not None:
            info_parts.append(f"Patterns Analysed: {pattern_count}")

        if info_parts:
            info_para = doc.add_paragraph()
            info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = info_para.add_run("  |  ".join(info_parts))
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        doc.add_page_break()

    # ------------------------------------------------------------------
    # Executive Summary
    # ------------------------------------------------------------------

    def _add_executive_summary(self, doc: Document, analysis: dict, analysis_json: dict) -> None:
        """Add the Executive Summary section with health score, breakdown table, and overview."""
        doc.add_heading("Executive Summary", level=1)

        # Health score prominent display
        health_score = analysis.get("health_score")
        if health_score is not None:
            score_para = doc.add_paragraph()
            run = score_para.add_run("Overall Health Score: ")
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

            score_run = score_para.add_run(f"{health_score}")
            score_run.bold = True
            score_run.font.size = Pt(14)
            score_color = self._score_color(health_score)
            score_run.font.color.rgb = score_color

            suffix_run = score_para.add_run(" / 100")
            suffix_run.font.size = Pt(14)
            suffix_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # Score breakdown table
        breakdown = analysis.get("score_breakdown_json", {})
        if isinstance(breakdown, dict) and breakdown:
            doc.add_paragraph("")  # spacer

            breakdown_heading = doc.add_paragraph()
            run = breakdown_heading.add_run("Score Breakdown")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

            table = doc.add_table(rows=1, cols=2)
            table.style = "Light Shading Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.LEFT

            hdr = table.rows[0].cells
            for i, label in enumerate(["Metric", "Score"]):
                hdr[i].text = label
                for paragraph in hdr[i].paragraphs:
                    for r in paragraph.runs:
                        r.bold = True

            for key, value in breakdown.items():
                row = table.add_row()
                row.cells[0].text = key.replace("_", " ").title()
                row.cells[1].text = f"{value}" if value is not None else "N/A"

            doc.add_paragraph("")  # spacer

        # Executive summary text (AI deep analysis key)
        summary = analysis_json.get("executive_summary", "") or analysis_json.get("overview", "")
        if summary:
            overview_heading = doc.add_paragraph()
            run = overview_heading.add_run("Overview")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

            self._add_structured_text(doc, summary)

    # ------------------------------------------------------------------
    # Assessment Scores (summary table for all 9 areas)
    # ------------------------------------------------------------------

    def _add_assessment_scores(self, doc: Document, analysis_json: dict) -> None:
        """Add a summary table of all 9 assessment area ratings."""
        # Only add section if at least one area exists
        has_any = any(analysis_json.get(key) for _, key in ANALYSIS_AREAS)
        if not has_any:
            return

        doc.add_heading("Assessment Scores", level=1)

        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Shading Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        hdr = table.rows[0].cells
        for i, label in enumerate(["Area", "Rating", "Findings", "Recommendations"]):
            hdr[i].text = label
            for paragraph in hdr[i].paragraphs:
                for r in paragraph.runs:
                    r.bold = True

        for area_name, area_key in ANALYSIS_AREAS:
            area_data = analysis_json.get(area_key, {})
            if not isinstance(area_data, dict):
                continue

            row = table.add_row()
            row.cells[0].text = area_name
            row.cells[1].text = area_data.get("rating", "N/A")

            # Count all findings (all list values except recommendation keys)
            finding_count = self._count_findings(area_data)
            row.cells[2].text = str(finding_count)

            # Count recommendations
            rec_count = self._count_recommendations(area_data)
            row.cells[3].text = str(rec_count)

        doc.add_paragraph("")  # spacer

    # ------------------------------------------------------------------
    # Detailed Assessments (one section per area)
    # ------------------------------------------------------------------

    def _add_detailed_assessments(self, doc: Document, analysis_json: dict) -> None:
        """Add detailed breakdown for each of the 9 assessment areas."""
        has_any = any(analysis_json.get(key) for _, key in ANALYSIS_AREAS)
        if not has_any:
            return

        doc.add_heading("Detailed Assessments", level=1)

        for area_name, area_key in ANALYSIS_AREAS:
            area_data = analysis_json.get(area_key, {})
            if not isinstance(area_data, dict):
                continue

            doc.add_heading(area_name, level=2)

            # Rating line
            rating = area_data.get("rating", "")
            if rating:
                rating_para = doc.add_paragraph()
                run = rating_para.add_run("Rating: ")
                run.bold = True
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

                rating_run = rating_para.add_run(rating)
                rating_run.bold = True
                rating_run.font.color.rgb = self._rating_color(rating)

            # Findings — all list/dict-list values except recommendations
            findings_rendered = False
            for fkey, fval in area_data.items():
                if fkey in ("rating",) or fkey in _RECOMMENDATION_KEYS:
                    continue
                if isinstance(fval, list) and fval:
                    if not findings_rendered:
                        findings_heading = doc.add_paragraph()
                        run = findings_heading.add_run("Findings")
                        run.bold = True
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)
                        findings_rendered = True

                    # Sub-label for the finding category
                    sublabel = doc.add_paragraph()
                    run = sublabel.add_run(fkey.replace("_", " ").title())
                    run.italic = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

                    for item in fval:
                        if isinstance(item, dict):
                            # Render dict items as key-value summary
                            parts = []
                            for dk, dv in item.items():
                                if isinstance(dv, list):
                                    parts.append(f"{dk}: {', '.join(str(x) for x in dv)}")
                                else:
                                    parts.append(f"{dk}: {dv}")
                            doc.add_paragraph(" | ".join(parts), style="List Bullet")
                        else:
                            doc.add_paragraph(str(item), style="List Bullet")

            # Recommendations
            recs = area_data.get("recommendations", []) or area_data.get("consolidation_suggestions", [])
            if isinstance(recs, list) and recs:
                recs_heading = doc.add_paragraph()
                run = recs_heading.add_run("Recommendations")
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x22, 0x88, 0x22)

                for rec in recs:
                    doc.add_paragraph(str(rec), style="List Bullet")

            doc.add_paragraph("")  # spacer between areas

    # ------------------------------------------------------------------
    # Maturity Roadmap
    # ------------------------------------------------------------------

    def _add_maturity_roadmap(self, doc: Document, analysis_json: dict) -> None:
        """Add the Maturity & Actionable Roadmap section."""
        roadmap = analysis_json.get("maturity_roadmap", {})
        if not isinstance(roadmap, dict) or not roadmap:
            return

        doc.add_heading("Maturity & Roadmap", level=1)

        # Overall maturity
        overall = roadmap.get("overall_maturity", "")
        if overall:
            maturity_para = doc.add_paragraph()
            run = maturity_para.add_run("Overall Maturity: ")
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

            maturity_run = maturity_para.add_run(overall)
            maturity_run.bold = True
            maturity_run.font.size = Pt(14)
            maturity_run.font.color.rgb = self._maturity_color(overall)

        # Area maturity table
        area_maturity = roadmap.get("area_maturity", {})
        if isinstance(area_maturity, dict) and area_maturity:
            doc.add_paragraph("")
            area_heading = doc.add_paragraph()
            run = area_heading.add_run("Area Maturity Levels")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

            table = doc.add_table(rows=1, cols=2)
            table.style = "Light Shading Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.LEFT

            hdr = table.rows[0].cells
            hdr[0].text = "Area"
            hdr[1].text = "Maturity"
            for cell in hdr:
                for paragraph in cell.paragraphs:
                    for r in paragraph.runs:
                        r.bold = True

            for area, level in area_maturity.items():
                row = table.add_row()
                row.cells[0].text = area.replace("_", " ").title()
                row.cells[1].text = str(level)

            doc.add_paragraph("")

        # Prioritized actions
        actions = roadmap.get("prioritized_actions", [])
        if isinstance(actions, list) and actions:
            actions_heading = doc.add_paragraph()
            run = actions_heading.add_run("Prioritized Actions")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

            for act in actions:
                if not isinstance(act, dict):
                    continue

                priority = act.get("priority", "")
                action_text = act.get("action", "")
                impact = act.get("impact", "")
                effort = act.get("effort", "")
                affected = act.get("affected_patterns", [])

                # Header line
                rec_para = doc.add_paragraph()
                header_run = rec_para.add_run(f"{priority}. {action_text}")
                header_run.bold = True
                header_run.font.size = Pt(12)
                header_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

                # Impact + effort
                if impact or effort:
                    meta_para = doc.add_paragraph()
                    if impact:
                        run = meta_para.add_run("Impact: ")
                        run.bold = True
                        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
                        impact_run = meta_para.add_run(impact)
                        impact_run.font.color.rgb = self._effort_color(impact)
                        impact_run.bold = True
                    if impact and effort:
                        meta_para.add_run("  |  ")
                    if effort:
                        run = meta_para.add_run("Effort: ")
                        run.bold = True
                        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
                        effort_run = meta_para.add_run(effort)
                        effort_run.font.color.rgb = self._effort_color(effort)
                        effort_run.bold = True

                # Affected patterns
                if isinstance(affected, list) and affected:
                    affected_para = doc.add_paragraph()
                    run = affected_para.add_run("Affected: ")
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                    affected_para.add_run(", ".join(str(p) for p in affected)).font.size = Pt(10)

                doc.add_paragraph("")  # spacer between actions

    # ------------------------------------------------------------------
    # Raw text fallback
    # ------------------------------------------------------------------

    def _add_raw_text_fallback(self, doc: Document, analysis_json) -> None:
        """Fallback when analysis_json contains raw_text instead of structured data."""
        doc.add_heading("Analysis", level=1)

        if isinstance(analysis_json, dict):
            raw_text = analysis_json.get("raw_text", "")
        elif isinstance(analysis_json, str):
            raw_text = analysis_json
        else:
            raw_text = str(analysis_json)

        if raw_text:
            self._add_structured_text(doc, raw_text)
        else:
            doc.add_paragraph("No analysis data available.")

    # ------------------------------------------------------------------
    # Page numbers
    # ------------------------------------------------------------------

    def _add_page_numbers(self, doc: Document) -> None:
        """Add centered page numbers to the footer."""
        section = doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run()
        fldChar1 = parse_xml(
            r'<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls('w'))
        )
        run._r.append(fldChar1)

        run2 = paragraph.add_run()
        instrText = parse_xml(
            r'<w:instrText {} xml:space="preserve"> PAGE </w:instrText>'.format(nsdecls('w'))
        )
        run2._r.append(instrText)

        run3 = paragraph.add_run()
        fldChar2 = parse_xml(
            r'<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls('w'))
        )
        run3._r.append(fldChar2)

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def _count_findings(self, area_data: dict) -> int:
        """Count all finding items in an area (all lists except recommendations)."""
        count = 0
        for key, val in area_data.items():
            if key in ("rating",) or key in _RECOMMENDATION_KEYS:
                continue
            if isinstance(val, list):
                count += len(val)
        return count

    def _count_recommendations(self, area_data: dict) -> int:
        """Count recommendation items in an area."""
        count = 0
        for key in _RECOMMENDATION_KEYS:
            val = area_data.get(key, [])
            if isinstance(val, list):
                count += len(val)
        return count

    def _add_structured_text(self, doc: Document, text: str) -> None:
        """Add text to the document, handling markdown lists, bold, and tables."""
        if not text:
            return

        lines = text.split("\n")
        in_table = False
        table_rows = []

        for line in lines:
            stripped = line.strip()
            is_table_row = stripped.startswith("|") and stripped.endswith("|") and len(stripped) > 2

            if is_table_row:
                # Skip separator rows (lines like |---|---|)
                cleaned = stripped.replace("|", "").replace("-", "").replace(":", "").replace(" ", "")
                if cleaned:
                    table_rows.append(stripped)
                in_table = True
            else:
                if in_table and table_rows:
                    self._flush_table_to_doc(doc, table_rows)
                    table_rows = []
                    in_table = False

                if stripped.startswith("- ") or stripped.startswith("* "):
                    item_text = stripped[2:]
                    self._apply_bold_runs_to_paragraph(doc, item_text, style="List Bullet")
                elif stripped.startswith("### "):
                    doc.add_heading(stripped[4:], level=3)
                elif stripped.startswith("## "):
                    doc.add_heading(stripped[3:], level=2)
                elif stripped.startswith("# "):
                    doc.add_heading(stripped[2:], level=1)
                elif stripped:
                    self._apply_bold_runs_to_paragraph(doc, stripped)

        # Flush remaining table
        if table_rows:
            self._flush_table_to_doc(doc, table_rows)

    def _apply_bold_runs_to_paragraph(self, doc: Document, text: str, style: str = None) -> None:
        """Create a paragraph with bold runs for **text** segments."""
        if style:
            para = doc.add_paragraph(style=style)
        else:
            para = doc.add_paragraph()

        # Split on bold markers: **bold text**
        parts = re.split(r'(\*\*.+?\*\*)', text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = para.add_run(part[2:-2])
                run.bold = True
            else:
                para.add_run(part)

    def _flush_table_to_doc(self, doc: Document, rows: list) -> None:
        """Convert markdown table rows into a Word table."""
        if not rows:
            return

        def parse_cells(row_str: str) -> list:
            cells = row_str.split("|")
            if cells and not cells[0].strip():
                cells = cells[1:]
            if cells and not cells[-1].strip():
                cells = cells[:-1]
            return [c.strip() for c in cells]

        header_cells = parse_cells(rows[0])
        col_count = len(header_cells)
        if col_count == 0:
            return

        table = doc.add_table(rows=1, cols=col_count)
        table.style = "Light Shading Accent 1"

        # Header row
        hdr = table.rows[0].cells
        for i, cell in enumerate(header_cells):
            if i < col_count:
                clean = re.sub(r'\*\*(.+?)\*\*', r'\1', cell)
                hdr[i].text = clean
                for paragraph in hdr[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

        # Data rows
        for row_str in rows[1:]:
            row_cells = parse_cells(row_str)
            row = table.add_row()
            for i, cell in enumerate(row_cells):
                if i < col_count:
                    clean = re.sub(r'\*\*(.+?)\*\*', r'\1', cell)
                    row.cells[i].text = clean

        doc.add_paragraph("")  # spacer after table

    def _score_color(self, score) -> RGBColor:
        """Return a colour based on the numeric health score (0-100)."""
        try:
            score = float(score)
        except (ValueError, TypeError):
            return RGBColor(0x88, 0x88, 0x88)

        if score >= 75:
            return RGBColor(0x22, 0x88, 0x22)  # green
        elif score >= 50:
            return RGBColor(0xDD, 0x99, 0x00)  # amber
        else:
            return RGBColor(0xCC, 0x33, 0x33)  # red

    def _rating_color(self, rating: str) -> RGBColor:
        """Return a colour based on a rating label (STRONG/ADEQUATE/WEAK etc)."""
        label = str(rating).upper()
        if label in ("STRONG", "LOW_RISK", "CLEAN"):
            return RGBColor(0x22, 0x88, 0x22)  # green
        elif label in ("ADEQUATE", "MODERATE_RISK", "SOME_OVERLAP"):
            return RGBColor(0xDD, 0x99, 0x00)  # amber
        elif label in ("WEAK", "HIGH_RISK", "SIGNIFICANT_OVERLAP"):
            return RGBColor(0xCC, 0x33, 0x33)  # red
        else:
            return RGBColor(0x88, 0x88, 0x88)  # grey

    def _maturity_color(self, maturity: str) -> RGBColor:
        """Return a colour based on maturity level."""
        label = str(maturity).upper()
        if label in ("MANAGED", "OPTIMIZING"):
            return RGBColor(0x22, 0x88, 0x22)  # green
        elif label in ("DEFINED", "DEVELOPING"):
            return RGBColor(0xDD, 0x99, 0x00)  # amber
        elif label == "INITIAL":
            return RGBColor(0xCC, 0x33, 0x33)  # red
        else:
            return RGBColor(0x88, 0x88, 0x88)

    def _effort_color(self, effort_label: str) -> RGBColor:
        """Return a colour based on a LOW/MEDIUM/HIGH label."""
        label = str(effort_label).upper()
        if label == "LOW":
            return RGBColor(0x22, 0x88, 0x22)  # green
        elif label == "MEDIUM":
            return RGBColor(0xDD, 0x99, 0x00)  # amber
        elif label == "HIGH":
            return RGBColor(0xCC, 0x33, 0x33)  # red
        else:
            return RGBColor(0x88, 0x88, 0x88)  # grey
